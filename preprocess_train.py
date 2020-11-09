import csv
import sys
import docx
from docx import Document
import io
import shutil
import os

from gensim.models.doc2vec import Doc2Vec, TaggedDocument
from nltk.tokenize import word_tokenize
import natsort

import pandas as pd
from sklearn.cluster import KMeans
import numpy as np
import matplotlib.pyplot as plt
from sklearn.preprocessing import StandardScaler
import pickle

from array import *
# import hierarchical clustering libraries
import scipy.cluster.hierarchy as sch
from sklearn.cluster import AgglomerativeClustering

import sys, json

class Preprocess:

    def read_in():
        lines = sys.stdin.readlines()
        # Since our input would only be having one line, parse our JSON data from that
        return json.loads(lines[0])

    def csvToDocx(filename):
        path= "docx/"

        with open(filename, newline='') as f:
            csv_reader = csv.reader(f) 

            csv_headers = next(csv_reader)
            csv_cols = len(csv_headers)
            i = 1
            for row in csv_reader:
                doc = docx.Document()
                formatted_output = str(row[1]).replace('\\n', '\n').replace('\\t', '\t')
                formatted_output.encode().decode()
                # formatted_output = unicodetoascii(formatted_output)
                formatted_output = formatted_output.replace('b\'', '')
                formatted_output = formatted_output.replace('b\"', '')
                formatted_output = formatted_output.replace('\'', '')
                formatted_output = formatted_output.replace('\"', '')
                formatted_output = formatted_output.replace(':', '')
                formatted_output = formatted_output.replace('\n\n', '\n')

                doc.add_paragraph(formatted_output)
                doc.add_page_break()
                doc.save("docx/"+str(row[0])+".docx")
                i = i+1

    def convertDocxToText(path):
        for d in os.listdir(path):
            fileExtension=d.split(".")[-1]
            if fileExtension =="docx":
                docxFilename = path + d
                # print(docxFilename)
                document = Document(docxFilename)
                textFilename = "txt/" + d.split(".")[0] + ".txt"
                with io.open(textFilename,"w", encoding="utf-8") as textFile:
                    for para in document.paragraphs: 
                        textFile.write(para.text)

    def docToVec(filenames):

        #Create empty list to store file contents
        doc = []

        for filename in natsort.natsorted(filenames,reverse=False):
            with open(os.path.join('./txt', filename)) as file:
                content = file.read()
                doc.append(content.lower())       

        # Tokenization of each document
        tokenized_doc = []
        for d in doc:
            tokenized_doc.append(word_tokenize(d.lower()))

        # Convert tokenized document into gensim formated tagged data
        tagged_data = [TaggedDocument(d, [i]) for i, d in enumerate(tokenized_doc)]
            
        ## Train doc2vec model
        model = Doc2Vec(tagged_data, vector_size=20, window=2, min_count=1, workers=4, epochs = 100)
        # Save trained doc2vec model
        model.save("sl_doc2vec.model")

    def vecToCsv(cia_model):
        # data to be written row-wise in csv fil 
        model= Doc2Vec.load(cia_model)
        data = model.docvecs 
        
        # opening the csv file in 'w+' mode 
        file = open('vecs.csv', 'w+', newline ='') 
        
        # writing the data into the file 
        with file:
            try:     
                write = csv.writer(file) 
                write.writerows(data)
            except KeyError:
                pass
    

    def cluster(vecs, txt_folder):
             
        df = pd.read_csv(vecs, header = None)
        col_size = df.shape[1]

        df.columns = ['vec'+str(x) for x in range(0,col_size)]
        df = df.applymap('{:.6f}'.format)

        df.to_csv("input_with_index.csv", header = True, index = True)

        x = pd.read_csv("input_with_index.csv")

        X = StandardScaler().fit_transform(x)

        new_row_size = X.shape[0]
        #new_row_size = row_size + 13*n (on first trials, the algorithm was clustering almost 13 inputs into one)
        print(new_row_size)
        n = round((new_row_size)/20)

        initial_clusters = 1
        n_clusters = initial_clusters

        if n == 0:
            n_clusters = n_clusters
        else:

            for l in range(1,10000):
    
                if n == l:
                    n_clusters = initial_clusters + l
                else:
                    n_clusters = n_clusters

            print(n_clusters)


        hc = AgglomerativeClustering(n_clusters=n_clusters, affinity = 'euclidean', linkage = 'ward')

        y_hc = hc.fit_predict(X)

        pickle.dump(hc, open(section+".pkl", "wb"))

data = []
datas = Preprocess.read_in()
for item in datas:
    data.append(item)
    
lines = data[0]
section = data[1]

a = Preprocess.csvToDocx(lines)
b = Preprocess.convertDocxToText('docx/')
filenames = os.listdir('./txt')
c = Preprocess.docToVec(filenames)
d = Preprocess.vecToCsv('sl_doc2vec.model')
x = Preprocess.cluster('vecs.csv','txt')

