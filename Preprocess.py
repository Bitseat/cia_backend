import csv
import sys
import docx
from docx import Document
import io
import shutil
import os

from gensim.models.doc2vec import Doc2Vec, TaggedDocument
from nltk.tokenize import word_tokenize
import natsort

import pandas as pd
from sklearn.cluster import KMeans
import numpy as np
import matplotlib.pyplot as plt
from sklearn.preprocessing import StandardScaler
import pickle
import pymongo 
import json

class Preprocess:
    def __init__(self):
        pass

    def append(self, issueid, reqtext, filename):
        
        fields = []

        fields.append(issueid)
        fields.append(reqtext)

        with open(filename, 'a') as f:
            writer = csv.writer(f)
            writer.writerow(fields)

    def csvToDocx(self,filename):
        path= "docx/"

        with open(filename, newline='') as f:
            csv_reader = csv.reader(f) 

            csv_headers = next(csv_reader)
            csv_cols = len(csv_headers)
            i = 1
            for row in csv_reader:
                doc = docx.Document()
                formatted_output = str(row[1]).replace('\\n', '\n').replace('\\t', '\t')
                formatted_output.encode().decode()
                # formatted_output = unicodetoascii(formatted_output)
                formatted_output = formatted_output.replace('b\'', '')
                formatted_output = formatted_output.replace('b\"', '')
                formatted_output = formatted_output.replace('\'', '')
                formatted_output = formatted_output.replace('\"', '')
                formatted_output = formatted_output.replace(':', '')
                formatted_output = formatted_output.replace('\n\n', '\n')

                doc.add_paragraph(formatted_output)
                doc.add_page_break()
                doc.save("docx/"+str(row[0])+".docx")
                i = i+1

    def convertDocxToText(self,path):
        for d in os.listdir(path):
            fileExtension=d.split(".")[-1]
            if fileExtension =="docx":
                docxFilename = path + d
                # print(docxFilename)
                document = Document(docxFilename)
                textFilename = "txt/" + d.split(".")[0] + ".txt"
                with io.open(textFilename,"w", encoding="utf-8") as textFile:
                    for para in document.paragraphs: 
                        textFile.write(para.text)

    def docToVec(self,filenames):

        #Create empty list to store file contents
        doc = []

        for filename in natsort.natsorted(filenames,reverse=False):
            with open(os.path.join('./txt', filename)) as file:
                content = file.read()
                doc.append(content.lower())       

        # Tokenization of each document
        tokenized_doc = []
        for d in doc:
            tokenized_doc.append(word_tokenize(d.lower()))

        # Convert tokenized document into gensim formated tagged data
        tagged_data = [TaggedDocument(d, [i]) for i, d in enumerate(tokenized_doc)]
            
        ## Train doc2vec model
        model = Doc2Vec(tagged_data, vector_size=20, window=2, min_count=1, workers=4, epochs = 100)
        # Save trained doc2vec model
        model.save("sl_doc2vec.model")

    def vecToCsv(self,cia_model):
        # data to be written row-wise in csv fil 
        model= Doc2Vec.load(cia_model)
        data = model.docvecs 
        
        # opening the csv file in 'w+' mode 
        file = open('vecs.csv', 'w+', newline ='') 
        
        # writing the data into the file 
        with file:
            try:     
                write = csv.writer(file) 
                write.writerows(data)
            except KeyError:
                pass

    def similarDocs(self,cia_model):
        ## Load saved doc2vec model
        model= Doc2Vec.load(cia_model)

        filenames = os.listdir('./txt')

        doc = []
        doc_name = []

        for filename in natsort.natsorted(filenames,reverse=False):
            with open(os.path.join('./txt', filename)) as file:
                content = file.read()
                doc.append(content.lower())
                filename = filename[:-4]
                doc_name.append(filename)

        # # find most similar doc 
        test_doc = word_tokenize(doc[-1].lower())
        # print(model.docvecs.most_similar(positive=[model.infer_vector(test_doc)],topn=5))

        print("finding similar requirements... \n")
        result = model.docvecs.most_similar(positive=[model.infer_vector(test_doc)],topn=5)
        for i in result:
            print(doc_name[i[0]]+'('+str(i[1])+')')
        return result

    def cluster(self,model_test, filename, vecs, txt_folder, candidates_fld):
        import shutil
        import os
        # test = pd.read_csv("testing.csv",header = None)
        model_test = pickle.load(open(model_test, "rb"))


        new = pd.read_csv(vecs, header = None)
        origin_data = pd.read_csv(filename)
        # new = origin.append(test) 
        new.to_csv("newone.csv")
        # origin[:1].to_csv("tt.csv")
        New = StandardScaler().fit_transform(new)

        pred = model_test.fit_predict(New)

        size = len(pred) -1
        d = pred[size]

        candidates = []
        for i in range(0, len(pred)):
            if pred[i] == d:
                candidates.append(i)

        issue_keys = []

        for index, row in origin_data.iterrows():
            for x in candidates:
                if index == x:
                    issue_keys.append(row[0])

        for i in issue_keys:
            shutil.copy(txt_folder + i+'.txt', candidates_fld)

        filenames = os.listdir('./candidates')
        doc = []

        for filename in natsort.natsorted(filenames,reverse=False):
            with open(os.path.join('./candidates', filename)) as file:
                content = file.read()
                doc.append(content.lower())       

        # Tokenization of each document
        tokenized_doc = []
        for d in doc:
            tokenized_doc.append(word_tokenize(d.lower()))

        # Convert tokenized document into gensim formated tagged data
        tagged_data = [TaggedDocument(d, [i]) for i, d in enumerate(tokenized_doc)]
            
        ## Train doc2vec model
        model = Doc2Vec(tagged_data, vector_size=20, window=2, min_count=1, workers=4, epochs = 100)
        # Save trained doc2vec model
        model.save("slc_doc2vec.model")

        ## Load saved doc2vec model
        model= Doc2Vec.load("slc_doc2vec.model")

        filenames = os.listdir('./candidates')
        filenames_txt = os.listdir('./txt')

        org_doc = []
        doc = []
        doc_name = []

        for filename in natsort.natsorted(filenames_txt,reverse=False):
            with open(os.path.join('./txt', filename)) as file:
                content = file.read()
                org_doc.append(content.lower())

        for filename in natsort.natsorted(filenames,reverse=False):
            with open(os.path.join('./candidates', filename)) as file:
                content = file.read()
                doc.append(content.lower())
                filename = filename[:-4]
                doc_name.append(filename)

        # # find most similar doc 
        test_doc = word_tokenize(org_doc[-1].lower())
        # print(model.docvecs.most_similar(positive=[model.infer_vector(test_doc)],topn=5))

        print("The top 5 similar requirements are: \n")
        result = model.docvecs.most_similar(positive=[model.infer_vector(test_doc)],topn=5)

        for i in result:
            print(doc_name[i[0]]+'('+str(i[1])+')')     

        import os, shutil
        folder = 'docx'
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))

        folder2 = 'txt'
        for filename in os.listdir(folder2):
            file_path = os.path.join(folder2, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))

        folder3 = 'candidates'
        for filename in os.listdir(folder3):
            file_path = os.path.join(folder3, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))

        return [result, doc_name]                

